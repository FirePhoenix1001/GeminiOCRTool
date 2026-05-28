import sys
import re
from docx import Document
from docx.oxml.ns import qn

def parse_sub_super(text, enable_super=True, enable_sub=True):
    r"""
    將文字拆分成普通文字、上標與下標區段，以匹配 Word 萬用字元：
    上標: ^{內容} (對應 [^94]\{([!\}]@)\})
    下標: _{內容} (對應 _\{([!\}]@)\})
    """
    segments = []
    i = 0
    n = len(text)
    
    while i < n:
        next_super = text.find('^{', i) if enable_super else -1
        next_sub = text.find('_{', i) if enable_sub else -1
        
        if next_super != -1 and (next_sub == -1 or next_super < next_sub):
            start_idx = next_super
            is_super = True
            marker_len = 2
        elif next_sub != -1 and (next_super == -1 or next_sub < next_super):
            start_idx = next_sub
            is_super = False
            marker_len = 2
        else:
            segments.append((text[i:], 'normal'))
            break
            
        if start_idx > i:
            segments.append((text[i:start_idx], 'normal'))
            
        close_idx = text.find('}', start_idx + marker_len)
        if close_idx != -1:
            content = text[start_idx + marker_len:close_idx]
            style = 'super' if is_super else 'sub'
            segments.append((content, style))
            i = close_idx + 1
        else:
            segments.append((text[start_idx:], 'normal'))
            break
            
    return segments

def changeWord(file_path, convert_minus_to_hyphen=False, convert_hyphen_to_minus=True, convert_super=False, convert_sub=False):
    doc = Document(file_path)
    math_functions = ['sin', 'cos', 'tan', 'log']
    
    for paragraph in doc.paragraphs:
        # 將原本 paragraph 裡的 Run 存成 list，避免一邊迭代一邊插入破壞結構
        for run in list(paragraph.runs):
            # 如果這個 Run 沒有文字 (例如是圖片、換行符等)，就直接跳過，這樣圖片就不會被刪掉
            if not run.text:
                continue
                
            old_text = run.text
            if convert_hyphen_to_minus:
                old_text = old_text.replace('-', '−')
            if convert_minus_to_hyphen:
                old_text = old_text.replace('−', '-')
                
            # 保存原本 Run 的基本格式
            original_bold = run.bold
            original_underline = run.underline
            
            # 進行上標與下標轉換解析
            segments = parse_sub_super(old_text, enable_super=convert_super, enable_sub=convert_sub)
            
            # 清空這個 run 原本的字串，我們接下來會依序塞入拆分好的 parts
            run.text = ""
            current_ref_run = run
            first_part = True
            
            for segment_text, style in segments:
                if not segment_text:
                    continue
                
                # 使用英文字母分組 (大小寫都抓)，以維持單字完整性
                parts = re.split(r'([a-zA-Z]+)', segment_text)
                
                for part in parts:
                    if not part:
                        continue
                        
                    if first_part:
                        target_run = run
                        first_part = False
                    else:
                        target_run = paragraph.add_run()
                        # 將新創建的 run 移動到上一個 target_run 的正後方，確保順序正確
                        current_ref_run._element.addnext(target_run._element)
                        target_run.bold = original_bold
                        target_run.underline = original_underline
                        current_ref_run = target_run
                    
                    target_run.text = part
                    target_run.font.name = 'Times New Roman'
                    target_run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    
                    # 設置上標或下標
                    if style == 'super':
                        target_run.font.superscript = True
                    elif style == 'sub':
                        target_run.font.subscript = True
                    else:
                        target_run.font.superscript = False
                        target_run.font.subscript = False
                    
                    # 只有 "完全由小寫英文字母組成" 的部分才進行斜體邏輯
                    if part.isalpha() and part.islower():
                        if part in math_functions:
                            target_run.italic = False  
                        else:
                            target_run.italic = True   
                    else:   
                        # 大寫或混合大小寫、數字、符號皆不斜體
                        target_run.italic = False

    try:
        doc.save(file_path)
    except PermissionError:
        print("\n" + "-" * 30)
        print(f"請關閉檔案{file_path}後，再重新執行一次程式。")
        print("-" * 30 + "\n")
        sys.exit(1)

def inputWord(text, file):
    doc = Document()
    math_functions = ['sin', 'cos', 'tan', 'log']
    for line in text.split('\n'):
        p = doc.add_paragraph()
        
        # inputWord 直接將字串切分放入新文件，沒有圖片流失的問題
        old_text = line.replace('-', '−')
        parts = re.split(r'([a-zA-Z]+)', old_text)
        
        for part in parts:
            if not part:
                continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            
            if part.isalpha() and part.islower():
                if part in math_functions:
                    run.italic = False  
                else:
                    run.italic = True   
            else:   
                run.italic = False
                
    doc.save(f"{file}.docx")   